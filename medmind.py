from llama_index.indices.managed.vectara import VectaraIndex
from dotenv import load_dotenv
import os
from PyPDF2 import PdfReader
from docx import Document
from sentence_transformers import SentenceTransformer
from Bio import Entrez
from llama_index.core.schema import Document
from llama_index.core import VectorStoreIndex, SimpleDirectoryReader, StorageContext
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.llms.together import TogetherLLM
from llama_index.core.llms import ChatMessage, MessageRole
from langchain.chains.question_answering import load_qa_chain
from langchain.text_splitter import CharacterTextSplitter
from langchain.embeddings.huggingface import HuggingFaceBgeEmbeddings
from langchain_community.vectorstores import Chroma 
from langchain.text_splitter import CharacterTextSplitter
import io
from transformers import AutoModelForSequenceClassification, AutoTokenizer
import streamlit as st
from googleapiclient.discovery import build
from typing import List, Optional
from llama_index.core import Settings

load_dotenv()

os.environ["VECTARA_INDEX_API_KEY"] = os.getenv("VECTARA_INDEX_API_KEY", "zwt_ni_bLu6MRQXzWKPIU__Uubvy_0Xz_FEr-2sfUg")
os.environ["VECTARA_QUERY_API_KEY"] = os.getenv("VECTARA_QUERY_API_KEY", "zwt_ni_bLu6MRQXzWKPIU__Uubvy_0Xz_FEr-2sfUg")
os.environ["VECTARA_API_KEY"] = os.getenv("VECTARA_API_KEY", "zut_ni_bLoa0I3AeNSjxeZ-UfECnm_9Xv5d4RVBAqw")
os.environ["VECTARA_CORPUS_ID"] = os.getenv("VECTARA_CORPUS_ID", "2")
os.environ["VECTARA_CUSTOMER_ID"] = os.getenv("VECTARA_CUSTOMER_ID", "2653936430")
os.environ["TOGETHER_API"] = os.getenv("TOGETHER_API", "7e6c200b7b36924bc1b4a5973859a20d2efa7180e9b5c977301173a6c099136b")
os.environ["GOOGLE_SEARCH_API_KEY"] = os.getenv("GOOGLE_SEARCH_API_KEY", "AIzaSyBnQwS5kPZGKuWj6sH1aBx5F5bZq0Q5jJk")
os.environ["PINECONE_API_KEY"] = os.getenv("PINECONE_API_KEY", "4523c180-39fd-4c48-99e8-88164df85b0a")

# Initialize the Vectara index
vectara_index = VectaraIndex()

endpoint = 'https://api.together.xyz/inference'

# Load the hallucination evaluation model
model_name = "vectara/hallucination_evaluation_model"
model = AutoModelForSequenceClassification.from_pretrained(model_name)
tokenizer = AutoTokenizer.from_pretrained(model_name)

def search_pubmed(query: str) -> Optional[List[str]]:
    """
    Searches PubMed for a given query and returns a list of formatted results 
    (or None if no results are found).
    """
    Entrez.email = "jayashbhardwaj3@gmail.com"  # Use environment variable for email

    try:
        handle = Entrez.esearch(db="pubmed", term=query, retmax=3)
        record = Entrez.read(handle)
        id_list = record["IdList"]

        if not id_list:
            return None

        handle = Entrez.efetch(db="pubmed", id=id_list, retmode="xml")
        articles = Entrez.read(handle)

        results = []
        for article in articles['PubmedArticle']:
            try:
                medline_citation = article['MedlineCitation']
                article_data = medline_citation['Article']
                title = article_data['ArticleTitle']
                abstract = article_data.get('Abstract', {}).get('AbstractText', [""])[0]

                result = f"**Title:** {title}\n**Abstract:** {abstract}\n"
                result += f"**Link:** https://pubmed.ncbi.nlm.nih.gov/{medline_citation['PMID']}\n\n"
                results.append(result)
            except KeyError as e:
                print(f"Error parsing article: {article}, Error: {e}")

        return results

    except Exception as e:
        print(f"Error accessing PubMed: {e}")
        return None

    except IOError as e:
        print(f"Error accessing PubMed: {e}")
        return None


def chat_with_pubmed(article_text, article_link):
    """
    Engages in a chat-like interaction with a PubMed article using TogetherLLM.
    """
    try:
        llm = TogetherLLM(model="QWEN/QWEN1.5-14B-CHAT", api_key=os.environ['TOGETHER_API'])
        messages = [
            ChatMessage(role=MessageRole.SYSTEM, content="You are a helpful AI assistant summarizing and answering questions about the following medical research article: " + article_link),
            ChatMessage(role=MessageRole.USER, content=article_text)
        ]
        response = llm.chat(messages)
        return str(response) if response else "I'm sorry, I couldn't generate a summary for this article."
    except Exception as e:
        print(f"Error in chat_with_pubmed: {e}")
        return "An error occurred while generating a summary."

def search_web(query: str, num_results: int = 3) -> Optional[List[str]]:
    """
    Searches the web using the Google Search API and returns a list of formatted results
    (or None if no results are found).
    """
    try:
        service = build("customsearch", "v1", developerKey=os.environ["GOOGLE_SEARCH_API_KEY"])

        # Execute the search request
        res = service.cse().list(q=query, cx="877170db56f5c4629", num=num_results).execute()

        if "items" not in res:
            return None

        results = []
        for item in res["items"]:
            title = item["title"]
            link = item["link"]
            snippet = item["snippet"]
            result = f"**Title:** {title}\n**Link:** {link}\n**Snippet:** {snippet}\n\n"
            results.append(result)

        return results

    except Exception as e:
        print(f"Error performing web search: {e}")
        return None

def extract_info_and_create_index(uploaded_file):
    try:
        # Extract text based on file type
        if uploaded_file.name.endswith(".pdf"):

            pdf_reader = PdfReader(uploaded_file)
            text = ""
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                try:
                    text += page.extract_text()
                except Exception as e:
                    print(f"Error extracting text from PDF page {page_num}: {e}")
        elif uploaded_file.name.endswith(".docx"):

            doc = Document(uploaded_file)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        else:  # Assuming .txt or other text-based format
            text = uploaded_file.getvalue().decode("utf-8")

        # Handle large files by chunking
        text_chunks = []
        max_chunk_size = 256  # Adjust as needed
        for i in range(0, len(text), max_chunk_size):
            text_chunks.append(text[i : i + max_chunk_size])

        # Create documents and embeddings for each chunk
        documents = []
        embed_model = HuggingFaceBgeEmbeddings(model_name="BAAI/bge-base-en")
        Settings.embed_model = embed_model
        for chunk in text_chunks:
            doc = Document(chunk)
            documents.append(doc)

        # Create a Chroma vector store index
        index = Chroma.from_documents(documents, embed_model)

        return index

    except Exception as e:
        print(f"Error processing uploaded file: {e}")
        return None
def chat_with_document(index, user_query):
    try:
        # Retrieve relevant documents and their text content based on the user query
        relevant_docs = index.similarity_search(user_query, k=3)

        # Combine text from relevant documents (handling chunking)
        document_text = ""
        for doc in relevant_docs:
            for chunk in doc.chunks:  # Access chunks within each document
                document_text += chunk.page_content + "\n\n"

        # Use TogetherAI's QWEN 1.5 14B model for generating the response
        together_llm = TogetherLLM(model="QWEN/QWEN1.5-14B-CHAT", api_key=os.environ['TOGETHER_API'])

        # Use Langchain's question-answering chain to generate a response
        qa_chain = load_qa_chain(llm=together_llm, chain_type="stuff")
        response = qa_chain.run(input_documents=relevant_docs, question=user_query)

        return response
    except Exception as e:
        print(f"Error in chat_with_document: {e}")
        return "An error occurred while processing your request."

def medmind_chatbot(user_input, index, chat_history=None):
    if chat_history is None:
        chat_history = []

    response_text = ""
    try:
        if "uploaded_index" in st.session_state and st.session_state["uploaded_index"] is not None:
            response_text = chat_with_document(st.session_state["uploaded_index"], user_input)
        else:
            # If no document is uploaded, proceed with Vectara, PubMed, and Web searches
            query_str = user_input
            response = vectara_index.as_query_engine().query(query_str)
            vectara_response = f"**MedMind Vectara Knowledge Base Response:**\n{response.response}"

            # PubMed Search and Chat
            pubmed_results = search_pubmed(user_input)
            pubmed_response = "**PubMed Articles (Chat & Summarize):**\n\n"
            if pubmed_results:
                for article_text in pubmed_results:
                    title, abstract, link = article_text.split("\n")[:3]
                    chat_summary = chat_with_pubmed(abstract, link)
                    pubmed_response += f"{title}\n{chat_summary}\n{link}\n\n"
            else:
                pubmed_response += "No relevant PubMed articles found.\n\n"

            # Web Search
            web_results = search_web(user_input)
            web_response = "**Web Search Results:**\n\n"
            if web_results:
                web_response += "\n".join(web_results)
            else:
                web_response += "No relevant web search results found.\n\n"

            # Combine responses from different sources
            response_text = vectara_response + "\n\n" + pubmed_response + "\n\n" + web_response

        # Hallucination Evaluation
        def vectara_hallucination_evaluation_model(text):
            inputs = tokenizer(text, return_tensors="pt")
            outputs = model(**inputs)
            hallucination_probability = outputs.logits[0][0].item()  
            return hallucination_probability

        # Hallucination Evaluation (applies to all responses)
        hallucination_score = vectara_hallucination_evaluation_model(response_text)
        HIGH_HALLUCINATION_THRESHOLD = 0.8
        if hallucination_score > HIGH_HALLUCINATION_THRESHOLD:
            response_text = "I'm still under development and learning. I cannot confidently answer this question yet."

    except Exception as e:
        response_text = f"An error occurred while processing your request: {e}"

    chat_history.append((user_input, response_text))
    return response_text, chat_history

def show_info_popup():
    with st.expander("How to use MedMind"):
        st.write("""
        **MedMind is an AI-powered chatbot designed to assist with medical information.**

        **Capabilities:**

        *   **Answers general medical questions:** MedMind utilizes a curated medical knowledge base to provide answers to a wide range of health-related inquiries.
        *   **Summarizes relevant research articles from PubMed:** The chatbot can retrieve and summarize research articles from the PubMed database, making complex scientific information more accessible.
        *   **Provides insights from a curated medical knowledge base:** Beyond simple answers, MedMind offers additional insights and context from its knowledge base to enhance understanding. 
        *   **Perform safe web searches related to your query:** The chatbot can perform web searches using the Google Search API, ensuring the safety and relevance of the results.

        **Limitations:**

        *   **Not a substitute for professional medical advice:** MedMind is not intended to replace professional medical diagnosis and treatment. Always consult a qualified healthcare provider for personalized medical advice.
        *   **General knowledge and educational purposes:** The information provided by MedMind is for general knowledge and educational purposes only and may not be exhaustive or specific to individual situations.
        *   **Under development:** MedMind is still under development and may occasionally provide inaccurate or incomplete information. It's important to critically evaluate responses and cross-reference with reliable sources.
        *   **Hallucination potential:** While MedMind employs a hallucination evaluation model to minimize the risk of generating fabricated information, there remains a possibility of encountering inaccurate responses, especially for complex or niche queries.

        **How to use:**

        1.  **Type your medical question in the text box.**
        2.  **MedMind will provide a comprehensive response combining information from various sources.** This may include insights from its knowledge base, summaries of relevant research articles, and safe web search results.
        3.  **You can continue the conversation by asking follow-up questions or providing additional context.** This helps MedMind refine its search and offer more tailored information.
        4.  **in case the Medmind doesn't show the output please check your internet connection or rerun the same command**
        5.  **user can either chat with the documents or with generate resposne from vectara + pubmed + web search**
        5.  **chat with document feature is still under development so it would be better to avoid using it for now**
        """)

# Initialize session state
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []

# Define function to display chat history with highlighted user input and chatbot response
def display_chat_history():
    for user_msg, bot_msg in st.session_state.chat_history:
        st.info(f"**You:** {user_msg}")
        st.success(f"**MedMind:** {bot_msg}")

# Define function to clear chat history
def clear_chat():
    st.session_state.chat_history = []
# Define main function
def main():
    """
    Main function for the MedMind Streamlit application. 
    Sets up the UI, handles user interactions, and generates responses.
    """

    # Streamlit Page Configuration
    st.set_page_config(page_title="MedMind Chatbot", layout="wide")

    # Custom Styles
    st.markdown(
        """
        <style>
        .css-18e3th9 {
            padding-top: 2rem;
            padding-right: 1rem;
            padding-bottom: 2rem;
            padding-left: 1rem;
        }
        .stButton>button {
            background-color: #4CAF50;
            color: white;
        }
        body {
            background-color: #F0FDF4;
            color: #333333;
        }
        .stMarkdown h1, .stMarkdown h2, .stMarkdown h3, .stMarkdown h4, .stMarkdown h5, .stMarkdown h6 {
            color: #388E3C;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    # Title and Introduction
    st.title("MedMind Chatbot")
    st.write("Ask your medical questions and get reliable information!")

    # Example Questions (Sidebar)
    example_questions = [
        "What are the symptoms of COVID-19?",
        "How can I manage my diabetes?",
        "What are the potential side effects of ibuprofen?",
        "What lifestyle changes can help prevent heart disease?"
    ]
    st.sidebar.header("Example Questions")
    for question in example_questions:
        st.sidebar.write(question)

    # File Uploader (Sidebar)
    st.sidebar.header("Upload Document")
    uploaded_file = st.sidebar.file_uploader("Choose a document", type=["txt", "pdf", "docx"])
    if uploaded_file is not None:
        st.session_state.uploaded_index = extract_info_and_create_index(uploaded_file)
        st.sidebar.success("Document indexed successfully!")

    # Output Container
    output_container = st.container()

    # User Input and Chat History
    input_container = st.container()
    with input_container:
        user_input = st.text_input("You: ", key="input_placeholder", placeholder="Type your medical question here...")
        new_chat_button = st.button("Clear Chat")
        if new_chat_button:
            st.session_state.chat_history = []  # Clear chat history

    # Initialize the Vectara index
    vectara_index = VectaraIndex()
    # Initialize the Chroma index
    index = extract_info_and_create_index(uploaded_file) if uploaded_file else None

    if user_input:
        response, st.session_state.chat_history = medmind_chatbot(user_input, vectara_index, st.session_state.chat_history)
        with output_container:
            display_chat_history()

    # Information Popup
    show_info_popup()

if __name__ == "__main__":
    main()